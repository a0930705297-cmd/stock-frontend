from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = '微軟正黑體'
style.font.size = Pt(11)

def sf(run, bold=False, size=11, color=None):
    run.font.name = '微軟正黑體'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=None):
    p = doc.add_heading('', level=level)
    run = p.add_run(text)
    run.font.name = '微軟正黑體'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2e, 0x75, 0xb6)

def para(text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sf(run, bold=bold, size=size, color=color)
    return p

def two_row_table(labels, values, value_colors=None):
    t = doc.add_table(rows=2, cols=len(labels))
    t.style = 'Table Grid'
    for i, (l, v) in enumerate(zip(labels, values)):
        r1 = t.rows[0].cells[i].paragraphs[0].add_run(l)
        sf(r1, bold=True, size=9)
        r2 = t.rows[1].cells[i].paragraphs[0].add_run(v)
        c = value_colors[i] if value_colors else None
        sf(r2, size=9, color=c)

# ── 標題 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('隔日沖候選操作分析報告')
sf(run, bold=True, size=20, color=(0x1a, 0x56, 0x76))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('分析日期：2026-04-29　｜　策略：隔日沖候選觀察')
sf(run2, size=11, color=(0x60, 0x60, 0x60))
doc.add_paragraph()

# ── 策略說明 ──
heading('策略框架說明', level=2)
para('本清單依「多頭回檔觀察」邏輯產生，操作三段式流程：', size=11)
for item in [
    '第一段（盤後）：多頭結構中，5MA下穿10MA，量縮無放量黑K，列入隔日觀察',
    '第二段（盤中09:30+）：守住昨日低點，量能持續收縮，考慮分批試單',
    '第三段（確認進場）：帶下影線紅K 或 5MA 重新上彎為較強確認訊號',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    sf(run, size=10)

para('⚠️ 停損原則：現價跌破昨日低點即出場，不猶豫。', bold=True, color=(0xC0,0x00,0x00), size=10)
doc.add_paragraph()

# ── 總覽表 ──
heading('前五名總覽', level=2)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
for i, h in enumerate(['排名','代號/股名','現價','關鍵優勢','主要風險']):
    run = t.rows[0].cells[i].paragraphs[0].add_run(h)
    sf(run, bold=True, size=10)

overview = [
    ('🥇 1', '6139 亞翔',    '$680',   '法人成本支撐+紅K收82%+縮量0.68x',   '距20MA較遠(7.48%)'),
    ('🥈 2', '4977 眾達-KY', '$213',   '成交額最大+法人成本支撐+紅K收82%',  '停損距離需確認'),
    ('🥉 3', '8996 高力',    '$1,175', '收盤位置100%（全榜最強）+縮量0.5x', '昨低距現價10%，停損需改用開盤低'),
    ('  4',  '1560 中砂',    '$530',   '紅K收85%+縮量0.6x+昨低距2.3%',     '成交額8.47億偏低'),
    ('  5',  '3030 德律',    '$327.5', '成交額20億最大+法人成本支撐',        '收盤位置60%+昨低距7.8%'),
]
for rank, name, price, adv, risk in overview:
    row = t.add_row()
    for i, text in enumerate([rank, name, price, adv, risk]):
        run = row.cells[i].paragraphs[0].add_run(text)
        sf(run, size=9)
doc.add_paragraph()

# ── 個股詳細分析 ──
stocks = [
    {
        'rank': '第一名',
        'code': '6139 亞翔',
        'price': '$680',
        'tech_labels': ['訊號','K棒','MA5','MA10','MA20','距20MA','量比'],
        'tech_vals':   ['等回測','紅K','697','688.2','632.7','7.48%','0.68x'],
        'info_labels': ['成交額','收盤位置','守昨低','法人成本'],
        'info_vals':   ['17.05億','82%','$657','$695.9（支撐，距成本-2.28%）'],
        'info_colors': [None, None, (0xC0,0x00,0x00), (0x00,0x80,0x00)],
        'entry': '守住 $657，09:30後量能持續收縮且無明顯賣壓，分批試單。出現帶下影線紅K或5MA上彎為更強確認。',
        'stop':  '$657（昨日低點），距現價約3.4%，停損空間合理，跌破即出。',
        'target':'短線目標沿20MA方向，5MA重站10MA可持有至隔日。',
        'note':  '外資成本$695.9，現價在成本下方2.28%，法人護盤意願強。成交額17億流動性充足。縮量0.68x加紅K收82%，為本次清單綜合條件最佳標的。',
        'risk':  '距20MA尚有7.48%，若大盤轉弱可能繼續回落，嚴守停損。',
    },
    {
        'rank': '第二名',
        'code': '4977 眾達-KY',
        'price': '$213',
        'tech_labels': ['訊號','K棒','MA5','MA10','MA20','距20MA','量比'],
        'tech_vals':   ['等回測','紅K','220.2','223.9','205.8','3.5%','0.76x'],
        'info_labels': ['成交額','收盤位置','守昨低','法人成本'],
        'info_vals':   ['18.84億','82%','請確認掃描結果','$216.9（支撐，距成本-1.79%）'],
        'info_colors': [None, None, (0xC0,0x00,0x00), (0x00,0x80,0x00)],
        'entry': '守住昨日低點，09:30後確認量縮，可分批試單。距20MA僅3.5%，是等回測中最接近好入場位的標的。',
        'stop':  '昨日低點（請確認掃描數字），跌破即出。',
        'target':'短線回測20MA附近，站回10MA可持有。',
        'note':  '成交額18.84億為全榜最大，流動性最佳。外資成本$216.9，現價在成本下方1.79%，法人成本支撐確立。紅K收82%品質良好。KY股注意匯率及外資動態。',
        'risk':  'KY股若外資連續賣超可能加速下跌，關注外資籌碼動向。',
    },
    {
        'rank': '第三名',
        'code': '8996 高力',
        'price': '$1,175',
        'tech_labels': ['訊號','K棒','MA5','MA10','MA20','距20MA','量比'],
        'tech_vals':   ['等回測','紅K','1155','1162.5','1061.2','10.72%','0.5x'],
        'info_labels': ['成交額','收盤位置','守昨低','外資成本'],
        'info_vals':   ['16.94億','100%（全榜最強）','$1055（距現價10.2%，停損需改用）','$1080.6（現價高於成本8.73%）'],
        'info_colors': [None, (0x00,0x80,0x00), (0xC0,0x60,0x00), None],
        'entry': '⚠️ 特殊操作：昨低$1055距現價達10.2%，不適合當停損。改用開盤後前15分鐘低點作為當日停損依據，確認撐住後09:30再考慮進場。',
        'stop':  '開盤後前15分鐘形成的低點（非昨低$1055）。高價股波動大，部位需相應調整。',
        'target':'短線看5MA/10MA方向，縮量持續可持有至隔日。',
        'note':  '收盤位置100%是全榜最強訊號，昨日收在最高點，買盤完全主導。量比0.5x縮量良好，成交額近17億。高價股每點價差較大，獲利空間與風險都相對高。',
        'risk':  '距20MA達10.72%，昨低停損不適用，操作難度高於其他標的，建議有經驗者再考慮。',
    },
    {
        'rank': '第四名',
        'code': '1560 中砂',
        'price': '$530',
        'tech_labels': ['訊號','K棒','MA5','MA10','MA20','距20MA','量比'],
        'tech_vals':   ['等回測','紅K','529.4','531.7','504.4','4.87%','0.6x'],
        'info_labels': ['成交額','收盤位置','守昨低','外資成本'],
        'info_vals':   ['8.47億','85%','$518（距現價2.3%）','$504.1（現價高於成本5.15%）'],
        'info_colors': [None, (0x00,0x80,0x00), (0xC0,0x00,0x00), None],
        'entry': '守住 $518，09:30後量能繼續收縮，可小量試單。收盤位置85%為等回測中K棒品質最好。',
        'stop':  '$518（昨日低點），距現價約2.3%，停損緊湊，跌破即出。',
        'target':'短線目標回測20MA區域。',
        'note':  '紅K收盤位置85%為本次前五名中K棒品質最強（僅次於高力100%），縮量0.6x良好，昨低距現價2.3%停損設置最合理。主要缺點是成交額8.47億偏低，進出留意滑價。',
        'risk':  '成交額偏低，大單進出可能推動價格，快進快出需注意滑價問題。',
    },
    {
        'rank': '第五名',
        'code': '3030 德律',
        'price': '$327.5',
        'tech_labels': ['訊號','K棒','MA5','MA10','MA20','距20MA','量比'],
        'tech_vals':   ['等回測','紅K','327.3','327.4','301','8.81%','0.75x'],
        'info_labels': ['成交額','收盤位置','守昨低','法人成本'],
        'info_vals':   ['20.11億（全榜最大）','60%','$302（距現價7.8%）','$322.2（支撐，距成本+1.64%）'],
        'info_colors': [None, None, (0xC0,0x60,0x00), (0x00,0x80,0x00)],
        'entry': '守住 $302，但昨低距現價達7.8%，停損空間偏大。建議等「假跌破回站」訊號或更明確縮量確認再進場，不宜開盤即追。',
        'stop':  '$302（昨日低點），距現價約7.8%，空間偏大，建議降低部位控制風險。',
        'target':'短線目標靠近20MA方向，但回測空間大，持有時間可能較長。',
        'note':  '成交額20.11億為全榜最大，流動性最佳。法人成本支撐確立。主要缺點是收盤位置60%偏中，昨低距現價7.8%使停損設置困難，較適合等更好的切入點。',
        'risk':  '昨低停損距離過大，若大盤走弱可能在觸及停損前已損失較多，需特別注意部位大小。',
    },
]

for s in stocks:
    heading(f"{s['rank']}　{s['code']}　現價 {s['price']}", level=1)

    candle_colors = [None] * len(s['tech_vals'])
    for i, v in enumerate(s['tech_vals']):
        if v == '紅K':
            candle_colors[i] = (0xC0, 0x00, 0x00)
        elif v == '黑K':
            candle_colors[i] = (0x00, 0x00, 0x00)

    two_row_table(s['tech_labels'], s['tech_vals'], candle_colors)
    doc.add_paragraph()
    two_row_table(s['info_labels'], s['info_vals'], s['info_colors'])
    doc.add_paragraph()

    for label, content, color in [
        ('📍 進場策略', s['entry'], (0x1a, 0x56, 0x76)),
        ('🛑 停損設置', s['stop'], (0xC0, 0x00, 0x00)),
        ('🎯 目標方向', s['target'], (0x00, 0x80, 0x00)),
        ('📊 備註',     s['note'],   None),
        ('⚠️ 主要風險', s['risk'],   (0xC0, 0x60, 0x00)),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{label}　')
        sf(r1, bold=True, size=11)
        r2 = p.add_run(content)
        sf(r2, size=10, color=color)
    doc.add_paragraph()

# ── 免責聲明 ──
doc.add_page_break()
heading('免責聲明', level=2)
para('本報告僅供學習與研究用途，不構成任何投資建議。股票投資具有風險，操作前請自行評估風險承受能力。所有分析基於技術指標與籌碼資料，無法保證未來走勢。', size=10, color=(0x60,0x60,0x60))

out = r'c:\Users\User\Desktop\stock PD\隔日沖操作分析_20260429.docx'
doc.save(out)
print('saved:', out)
